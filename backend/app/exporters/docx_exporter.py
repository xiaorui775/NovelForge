from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.exporters.base import BaseExporter, ExportResult


class DocxExporter(BaseExporter):
    @property
    def format_name(self) -> str:
        return "docx"

    @property
    def display_name(self) -> str:
        return "DOCX 文档"

    @property
    def file_extension(self) -> str:
        return ".docx"

    @property
    def media_type(self) -> str:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    async def export(self, project_data: dict) -> ExportResult:
        doc = Document()

        style = doc.styles["Normal"]
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5

        # Cover page
        for _ in range(6):
            doc.add_paragraph("")

        title = doc.add_heading(project_data["project_name"], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")

        if project_data.get("project_description"):
            desc = doc.add_paragraph(project_data["project_description"])
            desc.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # Table of Contents
        toc_heading = doc.add_heading("目 录", level=1)
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")

        for chapter in project_data["chapters"]:
            title_text = chapter["title"]
            word_count = chapter.get("word_count", 0)
            if word_count > 0:
                title_text += f"  ({word_count} 字)"
            p = doc.add_paragraph(title_text)
            p.paragraph_format.space_after = Pt(6)

        doc.add_page_break()

        # Chapters
        for chapter in project_data["chapters"]:
            doc.add_heading(chapter["title"], level=1)

            if chapter.get("content"):
                for paragraph in chapter["content"].split("\n"):
                    paragraph = paragraph.strip()
                    if paragraph:
                        p = doc.add_paragraph(paragraph)
                        p.paragraph_format.first_line_indent = Cm(0.75)
            else:
                doc.add_paragraph("（本章尚未生成）")

            doc.add_page_break()

        # Ending
        doc.add_paragraph("")
        doc.add_paragraph("")
        end = doc.add_paragraph(f"全书完 · 共 {project_data['total_words']} 字")
        end.alignment = WD_ALIGN_PARAGRAPH.CENTER

        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        content = buffer.getvalue()

        return ExportResult(
            filename=f"{project_data['project_name']}.docx",
            content=content,
            media_type=self.media_type,
        )
